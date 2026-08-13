#!/usr/bin/env python3
"""Generate the Chinese WD1.0 draft for intelligent 3D digital-human coding."""

import argparse
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
GRAY = "E7E6E6"
plt.rcParams["font.sans-serif"] = ["Noto Sans CJK SC", "Droid Sans Fallback", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(document, headers, rows, widths=None):
    value = document.add_table(rows=1, cols=len(headers))
    value.alignment = WD_TABLE_ALIGNMENT.CENTER
    value.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(value.rows[0].cells[index], header, True)
        set_cell_shading(value.rows[0].cells[index], LIGHT_BLUE)
    for row in rows:
        cells = value.add_row().cells
        for index, item in enumerate(row):
            set_cell_text(cells[index], item)
    if widths:
        for row in value.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return value


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, code, separate, end])


def heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def paragraph(document, text, bold_prefix=None):
    value = document.add_paragraph()
    value.paragraph_format.first_line_indent = Cm(0.74)
    value.paragraph_format.line_spacing = 1.35
    if bold_prefix and text.startswith(bold_prefix):
        run = value.add_run(bold_prefix)
        run.bold = True
        value.add_run(text[len(bold_prefix):])
    else:
        value.add_run(text)
    return value


def code_block(document, text):
    value = document.add_paragraph()
    value.style = document.styles["Normal"]
    value.paragraph_format.left_indent = Cm(0.8)
    value.paragraph_format.right_indent = Cm(0.8)
    value.paragraph_format.space_before = Pt(4)
    value.paragraph_format.space_after = Pt(6)
    set_cell_shading_like_paragraph(value, "F2F2F2")
    run = value.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(8)


def set_cell_shading_like_paragraph(paragraph_value, fill):
    p_pr = paragraph_value._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def new_page(document):
    document.add_page_break()


def make_architecture_figure(path):
    fig, ax = plt.subplots(figsize=(11.5, 4.2))
    ax.axis("off")
    boxes = [
        (0.02, 0.54, 0.15, 0.25, "编码端重建\n人体模型、高斯与网络"),
        (0.23, 0.67, 0.16, 0.19, "运动参数子流\n无损预测与熵编码"),
        (0.23, 0.40, 0.16, 0.19, "网络参数子流\n量化与区间编码"),
        (0.23, 0.13, 0.16, 0.19, "姿态特征图子流\n学习式视频编码"),
        (0.46, 0.40, 0.13, 0.25, "复用与封装\n序列、帧和校验"),
        (0.66, 0.40, 0.13, 0.25, "分层解码\n依赖与缓存"),
        (0.85, 0.40, 0.13, 0.25, "解码端重建\n驱动和渲染"),
    ]
    for x, y, w, h, label in boxes:
        ax.add_patch(plt.Rectangle((x, y), w, h, facecolor="#D9EAF7", edgecolor="#1F4E78", lw=1.5))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=10)
    arrows = [
        ((0.17, 0.67), (0.23, 0.765)), ((0.17, 0.66), (0.23, 0.495)),
        ((0.17, 0.64), (0.23, 0.225)), ((0.39, 0.765), (0.46, 0.59)),
        ((0.39, 0.495), (0.46, 0.525)), ((0.39, 0.225), (0.46, 0.46)),
        ((0.59, 0.525), (0.66, 0.525)), ((0.79, 0.525), (0.85, 0.525)),
    ]
    for begin, end in arrows:
        ax.annotate("", xy=end, xytext=begin, arrowprops=dict(arrowstyle="->", color="#1F4E78", lw=1.4))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def configure(document):
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.3)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for style_name, size, color in [("Title", 26, BLUE), ("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11, "000000")]:
        style = document.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("三维数字人智能编码标准（WD1.0）")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("— ")
    add_field(footer, "PAGE")
    footer.add_run(" —")


def generate(output, figure_dir):
    output = Path(output)
    figure_dir = Path(figure_dir)
    figure_dir.mkdir(parents=True, exist_ok=True)
    architecture = figure_dir / "digital_human_coding_architecture.png"
    make_architecture_figure(architecture)

    doc = Document()
    configure(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(65)
    run = title.add_run("数字视音频编解码技术标准化工作组")
    run.bold = True
    run.font.size = Pt(17)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("AVS 工作文档 · 2026年8月13日").font.size = Pt(11)
    doc.add_paragraph()
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.paragraph_format.space_before = Pt(55)
    run = title2.add_run("三维数字人智能编码标准")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    wd = doc.add_paragraph()
    wd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = wd.add_run("WD1.0")
    run.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()
    table(doc, ["项目", "内容"], [
        ("来源", "AVS VRU 组 / HGC-Avatar 项目"),
        ("标题", "三维数字人智能编码标准 WD1.0"),
        ("状态", "工作草案（供技术讨论与实验验证）"),
        ("范围", "可流式动态三维数字人的分层表示、封装、解码与渲染"),
    ], [4, 11])
    note = doc.add_paragraph("说明：本文件仿照《三维体视频智能编码标准》WD4.0 的体例编写，是 WD1.0 技术草案，不表示已发布的正式国家、行业或团体标准。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].italic = True

    new_page(doc)
    heading(doc, "目录", 1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    paragraph(doc, "提示：在 Microsoft Word 或 LibreOffice 中选择“更新目录”，即可刷新页码。")

    new_page(doc)
    heading(doc, "前言", 1)
    paragraph(doc, "本文件规定一种面向动态三维数字人的智能编码框架。框架以参数化人体模型、可学习外观网络和时变姿态特征图为分层表示，在编码端完成数字人重建及中间表示导出，在传输层分别形成运动参数子流、网络参数子流和姿态特征图子流，在解码端恢复各层并驱动三维高斯表示完成自由视点渲染。")
    paragraph(doc, "本文件按照 AVS 工作文档常用的语法、语义、解析、解码和渲染结构编写。除明确标注为资料性的条款外，使用“应”“不应”的条款为规范性要求，使用“宜”“可”的条款为建议或允许事项。")
    paragraph(doc, "WD1.0 聚焦 HGC-Avatar 所采用的分层压缩逻辑，保持三类内容可独立编解码，同时定义统一序列层和访问单元层，使实现可以通过一个流水线自动完成压缩，不要求使用者手工依次调用三个编码器。")
    heading(doc, "引言", 1)
    paragraph(doc, "动态三维数字人通常包含随时间变化的人体姿态、高维外观网络以及由姿态驱动的空间特征。若直接逐帧传输完整几何或网络状态，码率和初始化时延较高。本文件利用不同内容的时间相关性和更新频率进行分层：运动层支持逐帧低时延更新；Posemap 层利用视频编码器去除时空冗余；网络层采用标量量化和概率模型编码，通常只在初始化或模型更新时发送。")
    paragraph(doc, "三层在语义上独立、在解码依赖上有序：网络层提供外观生成函数，运动层提供骨骼驱动参数，Posemap 层提供时变姿态条件。实现可缓存不变层并仅更新必要子流，从而适应直播、点播、交互通信及边云协同渲染。")

    new_page(doc)
    heading(doc, "1 范围", 1)
    paragraph(doc, "本文件规定动态三维数字人智能编码的总体框架、比特流语法、语义、解析过程、各子流解码过程、数字人重建和渲染过程，以及档次、级别、一致性和错误处理要求。")
    paragraph(doc, "本文件适用于以参数化人体模型和神经网络/三维高斯为核心表示的可动画数字人，包括离线点播、实时通信、自由视点视频、扩展现实和数字内容制作。具体训练方法、采集设备、相机标定方法和显示终端不属于规范性范围。")
    heading(doc, "2 规范性引用文件", 1)
    table(doc, ["编号", "文件", "用途"], [
        ("ISO/IEC 14496-10", "Advanced video coding", "视频子流参考"),
        ("ISO/IEC 23090-5", "Visual volumetric video-based coding", "体积视频术语参考"),
        ("IEEE 754", "Floating-point arithmetic", "数值表示"),
        ("SMPL-X", "Expressive Body Capture: 3D Hands, Face, and Body", "人体参数模型参考"),
        ("3D Gaussian Splatting", "Real-time Radiance Field Rendering", "渲染表示参考"),
    ], [3, 8, 5])
    paragraph(doc, "凡是注日期的引用文件，仅该日期对应的版本适用；不注日期的引用文件，其最新版本（包括所有修改单）适用。实现采用其他视频编码器时，应通过本文件定义的 codec_identifier 和外部配置参数明确标识。")

    new_page(doc)
    heading(doc, "3 术语和定义", 1)
    definitions = [
        ("三维数字人（3D digital human）", "具有可驱动人体结构、时变姿态和可渲染外观的三维人物表示。"),
        ("访问单元（access unit）", "对应一个输出时刻、包含该时刻解码所需一个或多个数据单元的集合。"),
        ("运动参数（motion parameters）", "描述根节点平移、全局旋转、身体/手部/面部关节姿态及形状的参数集合。"),
        ("Posemap", "由姿态或蒙皮变换生成、以规则二维阵列承载的时变条件特征。"),
        ("网络参数层（network parameter layer）", "用于恢复外观生成网络权重、偏置和必要状态的编码层。"),
        ("基础层（base layer）", "形成最低可解码数字人所必需的结构及运动信息层。"),
        ("增强层（enhancement layer）", "改善外观、局部细节或渲染质量但可按能力选择解码的层。"),
        ("随机访问点（random access point）", "无需引用该点之前访问单元即可开始正确解码的位置。"),
        ("参考重建（reference reconstruction）", "同一训练检查点在未压缩中间表示条件下产生的编码端输出。"),
    ]
    for term, definition in definitions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(term + "：")
        r.bold = True
        p.add_run(definition)
    heading(doc, "4 缩略语", 1)
    table(doc, ["缩略语", "英文", "中文"], [
        ("AU", "Access Unit", "访问单元"), ("CRC", "Cyclic Redundancy Check", "循环冗余校验"),
        ("DCVC", "Deep Contextual Video Compression", "深度上下文视频压缩"), ("GOP", "Group of Pictures", "图像组"),
        ("PSNR", "Peak Signal-to-Noise Ratio", "峰值信噪比"), ("SSIM", "Structural Similarity", "结构相似性"),
        ("SMPL-X", "Skinned Multi-Person Linear Model eXpressive", "可表达参数化人体模型"),
        ("WD", "Working Draft", "工作草案"),
    ], [2.4, 7.5, 5.5])

    new_page(doc)
    heading(doc, "5 约定、函数和描述符", 1)
    heading(doc, "5.1 语法描述约定", 2)
    paragraph(doc, "语法表从左至右给出语法元素名称、条件和描述符。花括号表示语法结构体，方括号表示数组索引。除特别规定外，多字节整数按大端序写入。保留位应写为 0，解码器应忽略其值。")
    table(doc, ["描述符", "含义"], [
        ("u(n)", "n 位无符号整数"), ("i(n)", "n 位有符号二进制补码整数"),
        ("f(n)", "n 位固定模式"), ("b(8)", "按字节对齐的数据"),
        ("ue(v)", "无符号指数哥伦布码"), ("se(v)", "有符号指数哥伦布码"),
        ("leb128(v)", "无符号变长整数"), ("bytes(n)", "n 个字节的不透明载荷"),
    ], [3, 12])
    heading(doc, "5.2 数学函数", 2)
    table(doc, ["函数", "定义"], [
        ("Clip3(a,b,x)", "min(max(x,a),b)"), ("Round(x)", "取最接近整数，二分之一远离零"),
        ("Log2Ceil(x)", "不小于 log2(x) 的最小整数"), ("CRC32(data)", "按声明多项式计算 32 位校验值"),
        ("Q(x,Δ)", "Round(x/Δ)，其中 Δ 为量化步长"), ("DQ(k,Δ)", "k×Δ"),
    ], [4, 11])
    heading(doc, "5.3 字节对齐", 2)
    paragraph(doc, "每个数据单元头在字节边界开始。payload_size 按字节计数且不包含起始码和单元头。需要加密或外部传输封装时，加密边界应与数据单元边界一致，以保留随机访问和选择性解码能力。")

    new_page(doc)
    heading(doc, "6 总体架构和解码依赖", 1)
    doc.add_picture(str(architecture), width=Cm(16.2))
    caption = doc.add_paragraph("图 1 — 三维数字人分层智能编码总体架构")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph(doc, "编码端首先运行一次完整重建/测试流程，同时导出逐帧运动参数、网络检查点和 Posemap 帧及归一化范围。压缩模块分别处理三类中间表示并复用到统一序列。解码端恢复中间表示后，再运行同构测试/渲染流程，模拟远端数字人重建。")
    table(doc, ["层", "时间特性", "典型编码", "解码依赖", "建议更新频率"], [
        ("结构/运动层", "逐帧", "预测 + Huffman/无损熵编码", "人体模型", "每 AU"),
        ("网络参数层", "低频", "标量量化 + Range coder", "网络结构标识", "序列开始或模型更新"),
        ("Posemap 层", "逐帧", "帧归一化 + 视频编码", "范围旁信息、参考帧", "每 AU/GOP"),
        ("渲染元数据层", "低频", "定长/变长语法", "坐标系、相机约定", "按需"),
    ], [2.5, 2.4, 4.2, 3.1, 3.1])
    heading(doc, "6.1 一次流水线操作", 2)
    code_block(doc, "encoder reconstruction\n  -> export {motion, network, posemap, range}\n  -> encode each logical substream\n  -> multiplex and transmit\n  -> decode {motion, network, posemap}\n  -> decoder reconstruction and rendering\n  -> conformance / fidelity verification")
    paragraph(doc, "三个逻辑 codec 不意味着部署或使用时必须手工执行三个命令。符合性实现应提供统一入口，自动安排导出、压缩、解压、依赖检查和解码端重建；独立 codec 接口仅用于调试、替换算法和研究消融。")

    new_page(doc)
    heading(doc, "7 比特流结构", 1)
    heading(doc, "7.1 数据单元类型", 2)
    table(doc, ["unit_type", "名称", "起始码后缀", "作用"], [
        ("0", "DH_SEQUENCE_HEADER", "0xB0", "序列级能力、坐标和时间信息"),
        ("1", "DH_ACCESS_UNIT_HEADER", "0xB1", "帧序号、时间戳和随机访问标志"),
        ("2", "DH_STRUCTURE_UNIT", "0xB2", "人体拓扑、参数布局和绑定关系"),
        ("3", "DH_MOTION_UNIT", "0xB3", "逐帧人体姿态参数"),
        ("4", "DH_NETWORK_UNIT", "0xB4", "量化后的网络参数与概率模型信息"),
        ("5", "DH_POSEMAP_UNIT", "0xB5", "Posemap 视频载荷和范围旁信息"),
        ("6", "DH_METADATA_UNIT", "0xB6", "相机、质量、用户数据及扩展元数据"),
        ("7", "DH_SEQUENCE_END", "0xB7", "序列结束"),
    ], [2, 4.5, 3, 6])
    heading(doc, "7.2 公共单元头", 2)
    code_block(doc, "digital_human_unit() {\n  unit_start_code                 f(24)  // 0x000001\n  unit_type                       u(8)\n  unit_flags                      u(8)\n  layer_id                        u(6)\n  temporal_id                     u(2)\n  payload_size                    u(32)\n  payload_crc_present_flag        u(1)\n  reserved_zero_7bits             u(7)\n  payload                         bytes(payload_size)\n  if (payload_crc_present_flag) payload_crc32 u(32)\n}")
    paragraph(doc, "公共头允许接收端在不了解载荷内部语法时跳过未知单元。layer_id 为 0 表示基础层；更高 layer_id 表示增强层。网络单元和结构单元应出现在引用它们的运动/Posemap 单元之前，或由缓冲模型保证在呈现时刻前可用。")

    new_page(doc)
    heading(doc, "8 序列头及结构层语法与语义", 1)
    heading(doc, "8.1 序列头", 2)
    code_block(doc, "dh_sequence_header() {\n  dh_standard_version             u(8)\n  profile_idc                     u(8)\n  level_idc                       u(8)\n  progressive_sequence_flag       u(1)\n  low_latency_flag                u(1)\n  coordinate_system_idc           u(3)\n  length_unit_idc                 u(2)\n  reserved_zero_bit               u(1)\n  time_scale                      u(32)\n  num_units_in_tick               u(32)\n  frame_count_present_flag        u(1)\n  if (frame_count_present_flag) frame_count u(32)\n  human_id                        u(32)\n  model_schema_hash               bytes(32)\n}")
    table(doc, ["语法元素", "语义"], [
        ("dh_standard_version", "WD1.0 实验比特流应写 1；正式版本由标准机构分配。"),
        ("profile_idc", "指示必需工具集合，取值见附录 A。"),
        ("coordinate_system_idc", "0 右手 Y 向上；1 右手 Z 向上；其他值保留。"),
        ("length_unit_idc", "0 米；1 厘米；2 毫米。"),
        ("model_schema_hash", "网络结构、参数名称、形状和顺序规范化描述的 SHA-256。"),
    ], [5, 10])
    heading(doc, "8.2 结构单元", 2)
    paragraph(doc, "结构单元应声明人体模型类型、参数名称与维度、骨骼数量、关节父子关系、模板顶点/高斯基元数量以及网络结构标识。解码器不支持声明的结构时，应拒绝对应层，但可继续解析其他可独立层。")
    code_block(doc, "dh_structure_unit() {\n  body_model_idc                  u(8)\n  joint_count                     ue(v)\n  shape_parameter_count           ue(v)\n  expression_parameter_count      ue(v)\n  primitive_count                 ue(v)\n  parameter_layout_count          ue(v)\n  for (i=0; i<parameter_layout_count; i++) {\n    parameter_semantic_id         u(16)\n    component_count               ue(v)\n    numeric_format_idc            u(4)\n  }\n  external_asset_count            ue(v)\n  ...\n}")

    new_page(doc)
    heading(doc, "9 运动参数子流", 1)
    heading(doc, "9.1 参数集合", 2)
    paragraph(doc, "运动参数单元可包含 transl、global_orient、body_pose、left_hand_pose、right_hand_pose、jaw_pose、leye_pose、reye_pose、expression 和 betas。序列不使用的参数可通过 present_flag 省略。betas 等序列恒定参数宜只在随机访问点发送。")
    table(doc, ["参数", "典型维度", "预测方式", "规范性要求"], [
        ("transl", "3", "帧间差分", "应无损恢复"), ("global_orient", "3", "旋转向量差分", "应无损恢复"),
        ("body_pose", "63", "逐分量/关节预测", "应无损恢复或显式标记有损"),
        ("hand_pose", "2×45 或 PCA", "逐分量预测", "布局由结构单元声明"),
        ("expression", "实现声明", "帧内/帧间", "可选"), ("betas", "实现声明", "序列常量", "随机访问点可用"),
    ], [3, 2.4, 4, 5.5])
    heading(doc, "9.2 无损编码", 2)
    paragraph(doc, "WD1.0 参考模式以原始数组的数据类型、形状和字节序为输入，对时间轴执行可逆预测，将残差按符号映射后进行 Huffman 编码。码流应携带参数键、维度、数据类型、原始长度和校验摘要。解码后序列化内容的 SHA-256 应与编码端记录一致。")
    code_block(doc, "motion_parameter_block() {\n  parameter_semantic_id           u(16)\n  dtype_idc                       u(4)\n  prediction_mode                 u(3)\n  shape_dimension_count           u(3)\n  for (d=0; d<shape_dimension_count; d++) shape[d] ue(v)\n  symbol_count                    leb128(v)\n  huffman_table_length            leb128(v)\n  huffman_table                   bytes(huffman_table_length)\n  residual_payload_length         leb128(v)\n  residual_payload                bytes(residual_payload_length)\n}")
    heading(doc, "9.3 解码过程", 2)
    paragraph(doc, "解码器先验证参数布局和符号数量，再恢复 Huffman 表、解析残差、执行逆符号映射和逆预测，最后按声明的数据类型及形状组装数组。任何数组越界、符号数量不匹配或校验失败均应报告比特流错误，不应将未验证的数组交给人体模型。")

    new_page(doc)
    heading(doc, "10 网络参数子流", 1)
    heading(doc, "10.1 参数顺序和量化", 2)
    paragraph(doc, "网络参数应按照 model_schema_hash 所覆盖的规范顺序连接。BatchNorm 的 num_batches_tracked 等离散状态不参与连续权重量化，但应在解码检查点的 avatar_net 状态字典中恢复。")
    paragraph(doc, "对连续参数 x，编码器使用 k=Round(x/Δq) 量化，其中 Δq 由 q_index 查表得到。WD1.0 参考表包含 10 个以 10 为底在 10⁻⁷ 至 1 之间对数均匀分布的步长。q_index 增大表示量化步长增大、码率降低、失真通常上升。")
    table(doc, ["q_index", "量化步长 Δq（参考实现）", "用途"], [
        ("4", "约 1.2915×10⁻⁴", "高保真"), ("5", "约 7.7426×10⁻⁴", "高质量平衡点"),
        ("6", "约 4.6416×10⁻³", "中码率"), ("7", "约 2.7826×10⁻²", "低码率/预览"),
    ], [3, 6, 6])
    heading(doc, "10.2 概率模型和 Range coding", 2)
    paragraph(doc, "量化符号采用零均值对称概率模型。尺度参数由量化符号标准差估计，并映射到预定义正尺度查找表。尺度查找表应从 10⁻³ 到 10³ 单调覆盖；不得将自然对数值直接作为 logspace 的指数端点，以免产生无穷尺度。")
    code_block(doc, "dh_network_unit() {\n  network_id                       u(16)\n  model_schema_hash               bytes(32)\n  q_index                         u(16)\n  scale_index                     u(16)\n  arithmetic_symbol_max           u(32)\n  parameter_element_count         u(64)\n  payload_length                  u(64)\n  payload                         bytes(payload_length)\n}")
    heading(doc, "10.3 解码检查", 2)
    paragraph(doc, "解码器应按结构声明逐张量取出符号并反量化，验证元素总数、张量形状、参数名称和 schema 哈希。网络载入应使用严格模式；除规范明确允许的状态外，不应静默忽略缺失键或多余键。")

    new_page(doc)
    heading(doc, "11 Posemap 子流", 1)
    heading(doc, "11.1 帧表示与归一化", 2)
    paragraph(doc, "Posemap 原始特征可包含负数及超出显示范围的浮点值。编码前应逐帧或按组确定最小值/最大值，将特征可逆映射至视频编码器的整数样本范围；相应范围参数属于必需旁信息，应与帧建立一一对应关系并纳入总传输码率。")
    table(doc, ["字段", "含义", "要求"], [
        ("posemap_width/height", "二维承载图尺寸", "序列内变化时应发更新单元"),
        ("channel_count", "特征通道数", "通道打包规则应声明"),
        ("sample_bit_depth", "视频输入位深", "参考实现为 8 位"),
        ("range_min/range_max", "反归一化范围", "应可靠传输；可无损压缩"),
        ("codec_identifier", "视频 codec 及版本", "应可唯一解析对应载荷"),
        ("intra_period", "帧内刷新周期", "影响随机访问和错误传播"),
    ], [4, 6, 5])
    heading(doc, "11.2 视频编码", 2)
    paragraph(doc, "WD1.0 参考实现采用 DCVC 系列学习式视频编码器，将第一个帧及周期性刷新帧编码为 I 帧，其余编码为 P 帧。码率点由 rate_index 指定。其他符合性实现可使用传统或学习式视频编码器，但应在外层语法中携带 codec_identifier、模型摘要、码率点和关键帧位置。")
    code_block(doc, "dh_posemap_unit() {\n  posemap_id                       u(16)\n  codec_identifier                u(32)\n  codec_configuration_length      u(16)\n  codec_configuration             bytes(codec_configuration_length)\n  frame_index                     u(32)\n  random_access_flag              u(1)\n  range_side_information_flag     u(1)\n  reserved_zero_6bits             u(6)\n  if (range_side_information_flag) range_payload()\n  coded_posemap_length            u(32)\n  coded_posemap                   bytes(coded_posemap_length)\n}")
    heading(doc, "11.3 解码与同步", 2)
    paragraph(doc, "Posemap 解码帧数应与目标访问单元数匹配。解码器应按帧索引匹配范围旁信息，完成反归一化后再送入外观网络。缺失范围信息时，不应使用默认范围替代；可按错误恢复策略冻结上一有效 Posemap 或降级到基础层。")

    new_page(doc)
    heading(doc, "12 解析、缓冲和随机访问", 1)
    heading(doc, "12.1 解析过程", 2)
    code_block(doc, "while (!end_of_stream) {\n  locate_start_code();\n  parse_common_unit_header();\n  validate(payload_size, layer_id, temporal_id);\n  if (unit_type is supported) parse_payload();\n  else skip(payload_size);\n  if (payload_crc_present_flag) verify_crc();\n  dispatch_to_layer_buffer();\n}")
    paragraph(doc, "解析器应对 payload_size 设置实现上限，避免整数溢出和无界内存分配。遇到未知 unit_type 时，应按长度跳过并继续同步。若公共头损坏，可搜索下一个 0x000001 起始码，但恢复后的访问单元在获得新的随机访问点前可被标记为不完整。")
    heading(doc, "12.2 缓冲模型", 2)
    paragraph(doc, "网络层通常是大体积、低频数据，应允许分片传送并在 schema 校验完成后原子切换。运动层和 Posemap 层按帧序号进入抖动缓冲。呈现某访问单元前，解码器应确保所引用的结构、网络版本、运动帧和 Posemap 帧均可用。")
    table(doc, ["情形", "必需行为", "允许的降级"], [
        ("网络层尚未完成", "不得使用未完整网络", "继续使用上一网络版本"),
        ("运动帧丢失", "标记 AU 不完整", "保持上一姿态或插值"),
        ("Posemap P 帧丢失", "停止引用损坏参考链", "等待 I 帧/显示基础层"),
        ("CRC 错误", "丢弃对应单元并报告", "按层独立恢复"),
        ("时间戳乱序", "按解码顺序缓存、按呈现顺序输出", "超时丢弃"),
    ], [4, 6, 5])
    heading(doc, "12.3 随机访问", 2)
    paragraph(doc, "随机访问点应包含或可引用一个已完整缓存的结构/网络版本、运动绝对帧、Posemap I 帧和必要的范围旁信息。接收端从随机访问点开始时，不应依赖该点之前的运动差分或 Posemap 参考帧。")

    new_page(doc)
    heading(doc, "13 解码端数字人重建", 1)
    heading(doc, "13.1 初始化", 2)
    paragraph(doc, "解码器根据结构单元构建人体模型、规范姿态模板、三维高斯/神经外观网络及参数缓存。网络参数解码后应形成可严格载入的检查点。运动参数解码后应形成与训练/编码端键名、数据类型和维度一致的序列文件或等价内存对象。")
    heading(doc, "13.2 逐帧重建", 2)
    code_block(doc, "for each access unit t:\n  motion_t  = decode_motion(t)\n  posemap_t = decode_posemap(t, references)\n  body_t    = SMPLX(template, shape, motion_t)\n  features  = StyleUNet(posemap_t; decoded_network)\n  gaussians = deform(canonical_gaussians, body_t, features)\n  image_t   = gaussian_splat(gaussians, camera_t)\n  output(image_t, timestamp_t)")
    paragraph(doc, "坐标系、单位、旋转表示和关节顺序应与序列头及结构单元一致。任何隐式转换均应在实现文档中记录。对相同码流、相同渲染相机和相同确定性执行环境，解码端应产生可重复结果；浮点并行计算造成的微小差异可按附录一致性容限处理。")
    heading(doc, "13.3 编码端/解码端模拟", 2)
    paragraph(doc, "离线参考流程允许先执行一次编码端 test，保存中间结果；随后对中间结果压缩和解码，并以解码内容作为输入再次执行 test，以此模拟远端解码器。第二次 test 不应读取未压缩网络、原始运动文件或原始 Posemap 帧，否则该结果不构成完整的端到端压缩验证。")

    new_page(doc)
    heading(doc, "14 渲染过程", 1)
    heading(doc, "14.1 三维高斯属性", 2)
    paragraph(doc, "每个高斯基元至少具有中心位置、尺度、旋转、不透明度和颜色/球谐系数。属性可由规范空间中的持久参数与逐帧网络输出共同确定。人体蒙皮或非刚性形变将规范高斯映射到目标姿态。")
    table(doc, ["属性", "符号", "约束"], [
        ("中心", "μ∈R³", "变换后位于声明坐标系"), ("尺度", "s∈R⁺³", "各分量应为正且有限"),
        ("旋转", "r 或 q", "四元数应归一化"), ("不透明度", "α", "渲染前裁剪至有效范围"),
        ("颜色/SH", "c", "阶数和通道布局由结构声明"),
    ], [4, 4, 7])
    heading(doc, "14.2 投影和合成", 2)
    paragraph(doc, "渲染器按相机内外参将三维协方差投影到屏幕空间，以深度顺序执行透明度合成。为降低实现差异，渲染参数应声明背景色、近平面、远平面、颜色空间、输出分辨率和抗锯齿设置。")
    heading(doc, "14.3 质量测量", 2)
    paragraph(doc, "对训练/采集真值的评价可报告 PSNR、SSIM 和感知指标。仅比较解码端输出与未压缩编码端参考输出时，应称为“参考重建一致性”，不得等同于对真实图像的重建质量。码率应包含网络、运动、Posemap 视频码流及其范围旁信息、必要头部和校验数据。")

    new_page(doc)
    heading(doc, "15 档次和级别", 1)
    table(doc, ["profile_idc", "档次", "必需工具", "用途"], [
        ("1", "DH-Basic", "结构 + 无损运动 + 静态外观", "低复杂度驱动"),
        ("2", "DH-Posemap", "Basic + Posemap 视频层", "动态细节"),
        ("3", "DH-Hierarchical", "Posemap + 量化网络层 + 分层更新", "完整 HGC 流水线"),
        ("4", "DH-LowLatency", "Hierarchical + 低时延约束", "实时通信"),
    ], [2.3, 3.5, 6, 3.7])
    table(doc, ["level_idc", "最大帧率", "最大 Posemap 样本/帧", "最大网络载荷", "最大缓冲"], [
        ("10", "30", "512×512×3", "512 MiB", "1 GiB"),
        ("20", "60", "1024×1024×3", "1 GiB", "2 GiB"),
        ("30", "90", "2048×2048×3", "2 GiB", "4 GiB"),
        ("40", "120", "4096×4096×4", "4 GiB", "8 GiB"),
    ], [2.2, 2.5, 4.4, 3, 3])
    paragraph(doc, "上述级别约束为 WD1.0 建议值，用于实现协商和资源保护。后续工作草案可依据复杂度实验调整。解码器声明支持某档次/级别时，应支持该组合规定的全部必需语法和工具。")

    new_page(doc)
    heading(doc, "16 一致性、错误处理与安全", 1)
    heading(doc, "16.1 比特流一致性", 2)
    paragraph(doc, "符合本文件的比特流应满足语法范围、单元依赖、参考关系、元素数量和校验约束。编码器应生成至少一个可解码随机访问点。所有实际传输的旁信息均应计入码率统计。")
    heading(doc, "16.2 解码器一致性", 2)
    paragraph(doc, "解码器应对声明档次和级别内的符合性码流完成解析和重建，不发生越界访问、未定义张量载入或无限资源分配。运动参数无损模式应通过摘要一致性验证；网络层应报告量化误差和载入结果；Posemap 层应报告解码帧数和参考链状态。")
    heading(doc, "16.3 安全考虑", 2)
    table(doc, ["风险", "规范性缓解"], [
        ("超大 payload_size", "在分配内存前校验级别上限和剩余字节数"),
        ("恶意张量形状", "校验 schema、元素总数、乘法溢出和实现上限"),
        ("不可信模型执行", "比特流仅承载参数；网络结构来自受信 schema，禁止载入任意可执行对象"),
        ("路径注入", "语义标识不得直接作为文件系统路径；输出目录应隔离"),
        ("解压炸弹", "限制 Huffman 符号数、视频分辨率、帧数和缓冲量"),
        ("隐私泄漏", "对身份网络和运动数据实施访问控制、加密和留存策略"),
    ], [4, 11])
    heading(doc, "16.4 实验记录", 2)
    paragraph(doc, "实验报告应至少记录代码版本、环境、GPU、数据序列、帧范围、视角、所有编码参数、各子流字节数、旁信息字节数、运行时间、质量指标、异常和重试情况。报告应保存机器可读 JSON/CSV 结果以便审计。")

    new_page(doc)
    heading(doc, "附录 A（规范性）档次工具表", 1)
    table(doc, ["工具", "Basic", "Posemap", "Hierarchical", "LowLatency"], [
        ("序列/结构头", "必需", "必需", "必需", "必需"), ("无损运动层", "必需", "必需", "必需", "必需"),
        ("Posemap 视频层", "—", "必需", "必需", "必需"), ("量化网络层", "可选", "可选", "必需", "必需"),
        ("多层 layer_id", "—", "可选", "必需", "可选"), ("CRC", "可选", "可选", "建议", "建议"),
        ("低时延标志与缓冲约束", "—", "—", "可选", "必需"),
    ], [5, 2.3, 2.3, 2.6, 2.6])
    heading(doc, "附录 B（资料性）参考流水线", 1)
    code_block(doc, "# 单入口示例；各 codec 由流水线自动调用\npython -m hgc_avatar.pipeline --config <pipeline.yaml> --stage all\n\n# 批量分层实验\npython scripts/run_compression_matrix.py \\\n  --config configs/experiments/compression_matrix.server.yaml --jobs 4")
    paragraph(doc, "参考流水线先生成 encoder/render、encoder/posemap 和中间参数，然后产生 compression/network、compression/smpl、compression/posemap，最后生成 decoder/checkpoint、decoder/smpl_params.npz 和 decoder/render。verification.json 对比两次渲染。")
    heading(doc, "附录 C（资料性）码率核算示例", 1)
    code_block(doc, "total_bytes = network_header + network_payload\n            + motion_stream\n            + posemap_video_payload + posemap_range_side_info\n            + mandatory_container_headers + integrity_data\nKiB_per_frame = total_bytes / frame_count / 1024")
    paragraph(doc, "网络层为一次性初始化数据时，应同时报告“完整序列总字节数”和“按序列帧数摊销的 KiB/帧”。跨会话缓存网络的业务场景可另行报告稳态码率，但不得用稳态值替代完整传输成本。")

    new_page(doc)
    heading(doc, "附录 D（资料性）WD1.0 后续工作", 1)
    for item in [
        "统一 network schema 的规范化序列化规则及跨框架参数映射；",
        "补充 Posemap 通道打包、范围旁信息预测和熵编码规范；",
        "建立多厂商解码器的一致性码流、误差容限和测试向量；",
        "研究网络分块、渐进传输、局部更新及可伸缩增强层；",
        "形成实时档次的端到端时延、显存和功耗约束；",
        "补充身份保护、水印、加密和模型知识产权元数据。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    heading(doc, "参考文献", 1)
    references = [
        "[1] HGC-Avatar: Hierarchical Gaussian Compression for Streamable Dynamic 3D Avatars, ACM Multimedia 2025.",
        "[2] Kerbl B, et al. 3D Gaussian Splatting for Real-Time Radiance Field Rendering, ACM TOG, 2023.",
        "[3] Pavlakos G, et al. Expressive Body Capture: 3D Hands, Face, and Body from a Single Image, CVPR, 2019.",
        "[4] Li J, et al. Deep Contextual Video Compression, NeurIPS, 2021.",
        "[5] 《三维体视频智能编码标准》WD4.0，AVS N4120，2025.",
    ]
    for ref in references:
        doc.add_paragraph(ref)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "三维数字人智能编码标准 WD1.0"
    doc.core_properties.subject = "动态三维数字人的分层表示、编码、解码与渲染"
    doc.core_properties.author = "HGC-Avatar 项目 / AVS VRU 组工作草案"
    doc.core_properties.keywords = "三维数字人,HGC-Avatar,智能编码,WD1.0"
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", required=True)
    parser.add_argument("--figure-dir", default="docs/standards/figures")
    args = parser.parse_args()
    generate(args.output, args.figure_dir)
    print(args.output)


if __name__ == "__main__":
    main()
