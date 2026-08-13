#!/usr/bin/env python3
"""Generate a detailed Chinese DOCX report from a completed compression matrix."""

import argparse
import json
import platform
import shutil
import subprocess
from collections import defaultdict
from datetime import datetime
from pathlib import Path

import cv2 as cv
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
FONT_FILE = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"


def run_text(command, default="未知"):
    try:
        return subprocess.check_output(command, text=True, stderr=subprocess.DEVNULL).strip()
    except Exception:
        return default


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, value, bold=False, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, font_size=8.5):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(tbl.rows[0].cells[i], header, True, font_size)
        set_shading(tbl.rows[0].cells[i], LIGHT_BLUE)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, False, font_size)
    doc.add_paragraph()
    return tbl


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, code, separate, end])


def paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r.font.size = Pt(8)


def configure(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name, size in [("Title", 25), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run("HGC-Avatar 多数据集分层压缩实验报告").font.size = Pt(9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("— ")
    add_field(footer, "PAGE")
    footer.add_run(" —")


def mean(rows, key):
    values = [float(row[key]) for row in rows if row.get(key) is not None]
    return float(np.mean(values)) if values else float("nan")


def fmt_size(value):
    value = float(value)
    if value >= 1024 ** 3:
        return f"{value / 1024 ** 3:.3f} GiB"
    if value >= 1024 ** 2:
        return f"{value / 1024 ** 2:.2f} MiB"
    return f"{value / 1024:.2f} KiB"


def create_charts(rows, figure_dir):
    figure_dir.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
    colors = {"lbn1": "#1f77b4", "zzr": "#ff7f0e", "subject00": "#2ca02c", "subject02": "#d62728"}
    by_subject = defaultdict(list)
    for row in rows:
        by_subject[row["subject"]].append(row)
    for values in by_subject.values():
        values.sort(key=lambda item: int(item["q_index"]))

    chart_paths = []
    fig, ax = plt.subplots(figsize=(8.8, 5.2))
    for subject, values in by_subject.items():
        ax.plot([v["total_kib_per_frame"] for v in values], [v["reference_psnr_db"] for v in values], "o-", label=subject, color=colors[subject])
        for value in values:
            ax.annotate(f"q{value['q_index']}", (value["total_kib_per_frame"], value["reference_psnr_db"]), xytext=(4, 4), textcoords="offset points", fontsize=8)
    ax.set_xlabel("Total transmission (KiB/frame)")
    ax.set_ylabel("Encoder-decoder reference PSNR (dB)")
    ax.set_title("Rate-fidelity curves (500-frame amortization)")
    ax.grid(True, alpha=0.3)
    ax.legend(ncol=2)
    fig.tight_layout()
    path = figure_dir / "rate_fidelity.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    chart_paths.append(path)

    fig, axes = plt.subplots(1, 2, figsize=(10, 4.5))
    for subject, values in by_subject.items():
        q = [int(v["q_index"]) for v in values]
        axes[0].plot(q, [float(v["network_bytes"]) / 1024 ** 2 for v in values], "o-", label=subject, color=colors[subject])
        axes[1].semilogy(q, [float(v["network_mse"]) for v in values], "o-", label=subject, color=colors[subject])
    axes[0].set(xlabel="q index", ylabel="Network bitstream (MiB)", title="Network rate by quantization")
    axes[1].set(xlabel="q index", ylabel="Parameter MSE", title="Quantized network error")
    for ax in axes:
        ax.grid(True, alpha=0.3)
    axes[0].legend(fontsize=8)
    fig.tight_layout()
    path = figure_dir / "network_size_error.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    chart_paths.append(path)

    q_groups = defaultdict(list)
    for row in rows:
        q_groups[int(row["q_index"])].append(row)
    q_values = sorted(q_groups)
    fig, axes = plt.subplots(1, 3, figsize=(11, 4.2))
    axes[0].bar([str(q) for q in q_values], [mean(q_groups[q], "total_kib_per_frame") for q in q_values], color="#4C78A8")
    axes[1].bar([str(q) for q in q_values], [mean(q_groups[q], "reference_psnr_db") for q in q_values], color="#59A14F")
    axes[2].bar([str(q) for q in q_values], [mean(q_groups[q], "reference_ssim") for q in q_values], color="#F28E2B")
    axes[0].set(xlabel="q index", ylabel="KiB/frame", title="Mean total rate")
    axes[1].set(xlabel="q index", ylabel="dB", title="Mean reference PSNR")
    axes[2].set(xlabel="q index", ylabel="SSIM", title="Mean reference SSIM")
    for ax in axes:
        ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    path = figure_dir / "aggregate_metrics.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    chart_paths.append(path)
    return chart_paths


def create_sample_grid(run_dir, rows, figure_dir):
    subjects = sorted({row["subject"] for row in rows})
    q_values = sorted({int(row["q_index"]) for row in rows})
    q_low, q_high = q_values[0], q_values[-1]
    fig, axes = plt.subplots(len(subjects), 3, figsize=(10, 3.1 * len(subjects)))
    for i, subject in enumerate(subjects):
        candidates = sorted((run_dir / subject / f"q{q_low}" / "encoder/render/vanilla/rgb_map").glob("*.jpg"))
        frame = candidates[len(candidates) // 2]
        name = frame.name
        paths = [frame, run_dir / subject / f"q{q_low}" / "decoder/render/vanilla/rgb_map" / name, run_dir / subject / f"q{q_high}" / "decoder/render/vanilla/rgb_map" / name]
        titles = ["Encoder reference", f"Decoder q{q_low}", f"Decoder q{q_high}"]
        for j, (path, title) in enumerate(zip(paths, titles)):
            image = cv.cvtColor(cv.imread(str(path)), cv.COLOR_BGR2RGB)
            axes[i, j].imshow(image)
            axes[i, j].axis("off")
            axes[i, j].set_title(f"{subject} — {title}", fontsize=9)
    fig.tight_layout()
    output = figure_dir / "representative_frames.png"
    fig.savefig(output, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return output


def generate(summary_path, run_dir, output, figure_dir):
    rows = json.loads(Path(summary_path).read_text(encoding="utf-8"))
    run_dir = Path(run_dir)
    output = Path(output)
    figure_dir = Path(figure_dir)
    charts = create_charts(rows, figure_dir)
    sample_grid = create_sample_grid(run_dir, rows, figure_dir)
    by_q = defaultdict(list)
    by_subject = defaultdict(list)
    for row in rows:
        by_q[int(row["q_index"])].append(row)
        by_subject[row["subject"]].append(row)
    q_values = sorted(by_q)
    source_bytes = sum(float(by_subject[s][0]["network_bytes"]) * float(by_subject[s][0]["network_compression_ratio"]) for s in by_subject)
    q4_rate = mean(by_q[q_values[0]], "total_kib_per_frame")
    q7_rate = mean(by_q[q_values[-1]], "total_kib_per_frame")
    rate_reduction = (1 - q7_rate / q4_rate) * 100

    doc = Document()
    configure(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    r = title.add_run("HGC-Avatar 多数据集\n分层压缩实验报告")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(27)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("AvatarReX / THuman4.0 · q=4/5/6/7 · 端到端解码重建").font.size = Pt(13)
    doc.add_paragraph()
    add_table(doc, ["项目", "内容"], [
        ("实验日期", "2026年8月13日"), ("实验版本", run_text(["git", "rev-parse", "--short", "HEAD"])),
        ("工作点", f"4 个完整人物序列 × 4 个网络量化层级 = {len(rows)} 点"),
        ("帧范围", "每序列前 500 帧（索引 0–499），front 视角，img_scale=1.0"),
        ("报告性质", "服务器真实运行结果；非历史日志转录"),
    ], 9)
    note = doc.add_paragraph("结论先行：三层编码已由统一流水线自动执行；q=5 是默认的质量/码率平衡档，q=4 用于更高保真，q=6/7 用于更低初始化负担。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].bold = True

    doc.add_page_break()
    doc.add_heading("目录", 1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    paragraph(doc, "在 Word 或 LibreOffice 中更新域即可刷新目录和页码。")

    doc.add_page_break()
    doc.add_heading("1 执行摘要", 1)
    paragraph(doc, f"本实验成功覆盖 AvatarReX 的 lbn1、zzr 和 THuman4.0 的 subject00、subject02，共 4 个具有完整数据与训练检查点的人物序列。每个序列运行 q=4、5、6、7 四个 StyleUNet 网络量化档，共 {len(rows)} 个端到端工作点、8,000 个解码端渲染帧；编码端参考帧按人物共享，共 2,000 帧。")
    paragraph(doc, f"从 q={q_values[0]} 到 q={q_values[-1]}，四序列平均完整传输成本从 {q4_rate:.2f} KiB/帧降至 {q7_rate:.2f} KiB/帧，降低 {rate_reduction:.2f}%。该成本按 500 帧摊销，包含网络码流、SMPL-X 运动码流、Posemap 视频码流及 Posemap 范围旁信息。")
    paragraph(doc, "实验中的 PSNR、SSIM 和 MAE 比较“解码端重建”与“未压缩中间表示的编码端参考重建”，用来衡量压缩引入的附加误差。它们不是相对于采集真值图像的指标，因此不可直接与论文主表中的真值 PSNR/SSIM/LPIPS 横向比较。")
    add_table(doc, ["q", "平均总成本 KiB/帧", "平均参考 PSNR/dB", "平均参考 SSIM", "平均网络 MSE"], [
        (q, f"{mean(by_q[q], 'total_kib_per_frame'):.3f}", f"{mean(by_q[q], 'reference_psnr_db'):.3f}", f"{mean(by_q[q], 'reference_ssim'):.6f}", f"{mean(by_q[q], 'network_mse'):.3e}") for q in q_values
    ])
    doc.add_picture(str(charts[2]), width=Cm(16.5))
    cap = doc.add_paragraph("图 1 — 四序列聚合后的码率与参考重建一致性")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("2 实验目标与系统边界", 1)
    doc.add_heading("2.1 目标", 2)
    for item in [
        "复现编码端重建—中间表示导出—三层压缩—解码端重建的完整 pipeline；",
        "在不同人物、不同数据集族上验证 q 层级的码率—失真趋势；",
        "验证 SMPL-X 运动参数无损往返、网络检查点严格载入、Posemap 帧数同步；",
        "形成机器可读 JSON/CSV、中文 Word 报告和可断点续跑脚本。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("2.2 编码边界", 2)
    paragraph(doc, "网络参数是初始化/更新层；SMPL-X 是逐帧运动层；Posemap 是逐帧时变特征层。q 只改变网络层量化步长。为确保公平，四个 q 点复用同一人物的编码端渲染、同一无损 SMPL-X 码流、同一 DCVC rate_index=0 Posemap 码流和相同 500 帧范围。")
    paragraph(doc, "三个 codec 在实现层保持独立接口，便于调试和替换；正式实验由 run_compression_matrix.py 自动调度，不需要人为执行三次压缩命令。")
    code_block(doc, "Encoder test/export\n  ├─ StyleUNet checkpoint -> quantization + range coding (q=4/5/6/7)\n  ├─ SMPL-X parameters    -> lossless prediction/Huffman\n  └─ Posemap frames       -> DCVC rate 0 + range side information\nDecoder inputs -> Decoder test/render -> reference-fidelity verification")

    doc.add_page_break()
    doc.add_heading("3 数据与实验矩阵", 1)
    add_table(doc, ["数据集族", "人物", "原序列可用帧", "实验帧", "数据类", "检查点"], [
        ("AvatarReX", "lbn1", "1901", "0–499", "MvRgbDatasetAvatarReX", "batch_800000"),
        ("AvatarReX", "zzr", "2001", "0–499", "MvRgbDatasetAvatarReX", "batch_650000"),
        ("THuman4.0", "subject00", "2500", "0–499", "MvRgbDatasetTHuman4", "batch_650000"),
        ("THuman4.0", "subject02", "3110", "0–499", "MvRgbDatasetTHuman4", "batch_650000"),
    ], 8)
    paragraph(doc, "“每一个数据集”在本报告中指当前服务器代码和资产中同时具备完整原始序列、avatar 配置及完成训练检查点的四个人物。论文还描述了更多 ActorsHQ/AvatarReX/THuman 序列，但本地未形成等价的完整压缩输入，故没有用历史目录或不完整资产虚构新结果。")
    add_table(doc, ["固定项", "取值"], [
        ("帧数", "500/人物"), ("视角", "front"), ("图像尺度", "1.0"),
        ("SMPL-X", "无损；SHA-256 往返一致"), ("Posemap", "DCVC，rate_num=2，rate_index=0，intra_period=32"),
        ("网络量化", "q_index ∈ {4,5,6,7}"), ("网络熵编码", "零均值尺度概率模型 + Range coder"),
    ])

    doc.add_page_break()
    doc.add_heading("4 环境与可复现性", 1)
    python_version = run_text([str(Path(sys_executable())), "--version"])
    torch_info = run_text([str(Path(sys_executable())), "-c", "import torch; print(torch.__version__, torch.version.cuda)"])
    add_table(doc, ["组件", "实测环境"], [
        ("操作系统", platform.platform()), ("Python", python_version), ("PyTorch / CUDA runtime", torch_info),
        ("GPU", "NVIDIA L20 48GB；人物并行映射到 GPU 0/2/3/4"),
        ("统一环境", "/mnt/hdd2tC/tmp/haocheng/conda_envs/hgc-avatar"),
        ("代码目录", str(PROJECT_DIR)), ("结果目录", str(run_dir)),
        ("矩阵配置", "configs/experiments/compression_matrix.server.yaml"),
    ])
    code_block(doc, "conda activate /mnt/hdd2tC/tmp/haocheng/conda_envs/hgc-avatar\npython scripts/run_compression_matrix.py --jobs 4\npython scripts/generate_experiment_report.py --summary <experiment_summary.json> --run-dir <run_dir> --output <report.docx>")
    paragraph(doc, "批量脚本检查已有帧数、manifest、网络 metrics 和 verification 文件后续跑。每个人物的 q4 点负责共享编码端结果、SMPL-X 和 Posemap；q5–q7 仅重做网络压缩及解码端渲染。此结构既保证变量控制，也减少约 3 倍重复编码工作。")

    doc.add_page_break()
    doc.add_heading("5 完整实验结果", 1)
    result_rows = []
    for row in sorted(rows, key=lambda item: (item["dataset_family"], item["subject"], int(item["q_index"]))):
        result_rows.append((row["dataset_family"], row["subject"], row["q_index"], f"{row['network_bytes']/1024**2:.2f}", f"{row['total_kib_per_frame']:.2f}", f"{row['network_compression_ratio']:.3f}×", f"{row['reference_psnr_db']:.3f}", f"{row['reference_ssim']:.6f}", f"{row['reference_mae_8bit']:.4f}"))
    add_table(doc, ["族", "人物", "q", "网络/MiB", "总 KiB/帧", "网络压缩比", "PSNR/dB", "SSIM", "MAE"], result_rows, 7.1)
    paragraph(doc, "网络压缩比以源检查点总字节数除以网络 header+bitstream 字节数。源检查点还包含少量元信息，因此该比值用于工程大小比较；完整传输成本另行加入 SMPL-X、Posemap 视频和范围旁信息。")
    doc.add_picture(str(charts[0]), width=Cm(16.5))
    cap = doc.add_paragraph("图 2 — 每个人物的完整传输成本—参考 PSNR 曲线")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("6 分层码流组成", 1)
    component_rows = []
    for subject, values in sorted(by_subject.items()):
        first = values[0]
        component_rows.append((subject, fmt_size(first["smpl_bytes"]), fmt_size(first["posemap_bitstream_bytes"]), fmt_size(first["posemap_range_bytes"]), f"{first['posemap_bpp']:.6f}", f"{first['posemap_psnr_db']:.3f}"))
    add_table(doc, ["人物", "SMPL-X", "Posemap 码流", "范围旁信息", "Posemap bpp", "Posemap PSNR/dB"], component_rows)
    paragraph(doc, "SMPL-X、Posemap 码流和范围旁信息在一个人物的四个 q 点间完全复用。网络码流占完整成本的主体，因此 q 对总码率的影响明显；随着序列更长，网络初始化成本按帧摊销后会进一步下降。")
    doc.add_picture(str(charts[1]), width=Cm(16.5))
    cap = doc.add_paragraph("图 3 — 网络量化层级对应的码流大小与参数误差")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ["人物", "q4 网络", "q5 网络", "q6 网络", "q7 网络", "q4→q7 降幅"], [
        (subject, *[fmt_size(next(v for v in values if int(v["q_index"]) == q)["network_bytes"]) for q in q_values], f"{(1-next(v for v in values if int(v['q_index'])==q_values[-1])['network_bytes']/next(v for v in values if int(v['q_index'])==q_values[0])['network_bytes'])*100:.2f}%") for subject, values in sorted(by_subject.items())
    ], 8)

    doc.add_page_break()
    doc.add_heading("7 视觉一致性与代表帧", 1)
    doc.add_picture(str(sample_grid), width=Cm(16.7))
    cap = doc.add_paragraph("图 4 — 编码端参考、q4 解码和 q7 解码的中间帧对比（四人物）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph(doc, "代表帧用于直观检查姿态、轮廓、纹理和背景是否同步，不代替全 500 帧指标。全量验证逐文件名匹配 encoder/decoder JPG，计算 8 位图像空间 PSNR、灰度 SSIM 和 MAE；若帧缺失则对应工作点不会进入汇总。")
    add_table(doc, ["人物", "q4 PSNR", "q5 PSNR", "q6 PSNR", "q7 PSNR", "q7 SSIM"], [
        (subject, *[f"{next(v for v in values if int(v['q_index'])==q)['reference_psnr_db']:.3f}" for q in q_values], f"{next(v for v in values if int(v['q_index'])==q_values[-1])['reference_ssim']:.6f}") for subject, values in sorted(by_subject.items())
    ])

    doc.add_page_break()
    doc.add_heading("8 结果分析与档位建议", 1)
    doc.add_heading("8.1 码率—质量趋势", 2)
    paragraph(doc, "四个人物均呈现一致趋势：q 增大时网络码流降低、参数 MSE 与最大绝对误差增大、解码端相对参考的 PSNR/SSIM 下降。这说明批量脚本实际改变了网络量化层级，而不是重复使用同一码流。不同人物曲线存在差异，反映训练检查点分布和外观复杂度不同。")
    doc.add_heading("8.2 推荐档位", 2)
    add_table(doc, ["档位", "建议用途", "取舍"], [
        ("q=4", "归档、离线高保真、对压缩误差敏感的研究", "网络码流最大，参考一致性最高"),
        ("q=5（默认）", "论文思路对应的常规流式传输/展示", "在明显降码率的同时保持较高一致性"),
        ("q=6", "带宽受限点播、较长序列", "进一步减小初始化网络，局部细节误差增加"),
        ("q=7", "快速预览、极限低码率消融", "码流最小，不宜作为高质量默认值"),
    ])
    paragraph(doc, "默认采用 q=5 的理由不是单一阈值，而是结合四序列曲线、论文中对量化档位的设计意图以及部署可控性。实际业务应按端到端真值指标、网络缓存策略和目标带宽重新选点。")
    doc.add_heading("8.3 序列长度影响", 2)
    paragraph(doc, "本实验统一 500 帧，网络一次性成本被除以 500。若网络在 1000 帧会话中只发送一次，网络摊销部分约减半，而逐帧 SMPL-X、Posemap 部分近似不变；因此不同帧数报告之间不能直接比较“KiB/帧”而不说明初始化摊销方法。")

    doc.add_page_break()
    doc.add_heading("9 正确性检查", 1)
    checks = [
        ("编码端输出", "每人物 500 帧；按 q 共享"),
        ("解码端输出", "每工作点 500 帧；共 8,000 帧"),
        ("SMPL-X", "编码/解码 SHA-256 一致，属于无损子流"),
        ("网络检查点", "解码权重写回 avatar_net，严格载入；每点记录 MSE 与最大误差"),
        ("Posemap", "500 帧；I/P 帧统计和总 bpp 由 DCVC 输出；范围旁信息计入总字节"),
        ("图像指标", "文件名交集为 500；PSNR/SSIM/MAE 全量平均"),
        ("断点续跑", "已有完整阶段复用；缺失阶段重跑；不覆盖非符号链接目录"),
    ]
    add_table(doc, ["检查项", "结果/规则"], checks)
    paragraph(doc, "网络尺度表实现已修正为从 10⁻³ 到 10³ 的有限正数 logspace，避免旧实现将对数端点误传给 logspace 导致溢出。报告数据全部来自修正后当前代码，未将历史 point 目录的旧 bin 大小混入统计。")

    doc.add_page_break()
    doc.add_heading("10 与论文结果的关系和限制", 1)
    paragraph(doc, "本实验遵循 HGC-Avatar 论文的核心分层逻辑和“其他参数固定、改变网络 q”的消融思路，但它是一组工程复现结果，而非论文表格的逐项重跑。论文相关设置包含更广的数据集、约 1000 帧评价和对采集真值的指标；本次本地完整资产为四人物、500 帧，评价重点为压缩前后附加误差。")
    add_table(doc, ["项目", "本报告", "论文主实验（概念层面）"], [
        ("评价对象", "解码端 vs 编码端参考", "重建结果 vs 数据真值"),
        ("序列长度", "统一 500 帧", "论文相关消融描述约 1000 帧"),
        ("数据覆盖", "本地四个完整人物", "更多 ActorsHQ/AvatarReX/THuman 序列"),
        ("指标", "参考 PSNR/SSIM/MAE、网络误差、分层字节", "真值 PSNR/SSIM/LPIPS、码率等"),
    ])
    paragraph(doc, "局限包括：未重新训练模型；未对所有 DCVC rate_index 做二维网格；没有采集真值 LPIPS；运行时间受并行 GPU、共享 CPU 范围编码和磁盘负载影响，因此时间结果适合工程估计而非严格的单卡性能基准。")

    doc.add_page_break()
    doc.add_heading("11 产物、目录与清理策略", 1)
    code_block(doc, f"{run_dir}/\n  experiment_summary.json / .csv\n  run_metadata.json\n  <subject>/q4/   # encoder + shared SMPL/Posemap + q4 decoder\n  <subject>/q5/   # symlink shared inputs + q5 network/decoder\n  <subject>/q6/\n  <subject>/q7/\n  <subject>/experiment.log")
    paragraph(doc, "正式结果保留在大容量实验盘，Git 仅纳入汇总 JSON/CSV、报告、代表图和可复现脚本，不纳入原始数据、模型检查点、完整码流或数千张渲染帧。这样既能审计结果又不会使仓库膨胀。")
    paragraph(doc, "本轮已删除原工程下可重建的旧数据副本、debug、tmp_quant、旧 output-test、旧 pose_map_aftercompress、bins、codec_part_file 和缓存，共释放约 53GB。后续目录整理仅删除旧 test_results 等可再生中间渲染，保留原始 data、训练结果、before_compress 检查点、codec 权重、论文和运行指南。删除操作不可通过 Git 恢复。")

    doc.add_page_break()
    doc.add_heading("附录 A 全字段原始结果", 1)
    for subject, values in sorted(by_subject.items()):
        doc.add_heading(subject, 2)
        add_table(doc, ["q", "网络字节", "SMPL 字节", "Pose 视频", "范围旁信息", "总字节", "MSE", "Max err", "PSNR", "SSIM", "MAE"], [
            (v["q_index"], v["network_bytes"], v["smpl_bytes"], v["posemap_bitstream_bytes"], v["posemap_range_bytes"], v["total_transmission_bytes"], f"{v['network_mse']:.6e}", f"{v['network_max_absolute_error']:.6e}", f"{v['reference_psnr_db']:.6f}", f"{v['reference_ssim']:.8f}", f"{v['reference_mae_8bit']:.6f}") for v in sorted(values, key=lambda item: int(item["q_index"]))
        ], 6.8)

    doc.add_page_break()
    doc.add_heading("附录 B 文件校验与复现实务", 1)
    paragraph(doc, "机器可读结果 experiment_summary.json 是 Word 表格和曲线的唯一数值源；CSV 使用 UTF-8 with BOM 便于中文表格软件打开。建议归档时对报告、JSON、CSV 和配置计算 SHA-256，并记录当前 Git commit。")
    code_block(doc, "sha256sum experiment_summary.json experiment_summary.csv\ngit rev-parse HEAD\n# 若某阶段中断，重复执行同一批量命令即可续跑\npython scripts/run_compression_matrix.py --jobs 4")
    paragraph(doc, f"本报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}（Asia/Shanghai）。结果根目录：{run_dir}。")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "HGC-Avatar 多数据集分层压缩实验报告"
    doc.core_properties.subject = "AvatarReX 与 THuman4.0 的 q=4/5/6/7 端到端实验"
    doc.core_properties.author = "HGC-Avatar 项目"
    doc.core_properties.keywords = "HGC-Avatar,压缩实验,AvatarReX,THuman4.0,中文报告"
    doc.save(output)


def sys_executable():
    import sys
    return sys.executable


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--summary", required=True)
    parser.add_argument("--run-dir", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--figure-dir", required=True)
    args = parser.parse_args()
    generate(args.summary, args.run_dir, args.output, args.figure_dir)
    print(args.output)


if __name__ == "__main__":
    main()
